import os
import socket
from flask import Flask, render_template_string, url_for, redirect, make_response, request
import io
import base64
import sqlite3
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Mm
import subprocess
from jinja2 import Environment, FileSystemLoader
# defer importing docx2pdf until runtime to provide clearer errors if missing

app = Flask(__name__)
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'output')
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')
DB_PATH = Path(__file__).with_name('students.db')

# Configure wkhtmltopdf path; update if installed elsewhere
# We now use docx + docx2pdf for generation on demand.

env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))

QR_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Hall Ticket Download QR</title>
  <style>
    body { font-family: Arial, sans-serif; margin: 30px; }
    .wrap { max-width: 600px; margin: 0 auto; text-align: center; }
    img { width: 280px; height: 280px; }
    .hint { margin-top: 14px; color: #444; }
    code { background: #f4f4f4; padding: 2px 6px; }
  </style>
</head>
<body>
  <div class="wrap">
    <h2>Scan to Download Hall Ticket</h2>
    <p>Point your phone camera at the QR below.</p>
    <img src="data:image/png;base64,{{ qr_b64 }}" alt="Download QR">
    <div class="hint">
      If scanning on the same laptop, open <code>{{ download_url }}</code> in a browser.
    </div>
    <p><a href="{{ download_url }}">Direct download link</a></p>
  </div>
</body>
</html>
"""


def get_local_ip() -> str:
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
    except Exception:
        ip = '127.0.0.1'
    finally:
        s.close()
    return ip

INDEX_HTML = """
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8" />
  <title>Hall Ticket - Enter Register Number</title>
  <style>
    body { font-family: Arial, sans-serif; margin: 30px; }
    .wrap { max-width: 520px; margin: 0 auto; }
    form { display: flex; gap: 10px; }
    input { flex: 1; padding: 8px; }
    button { padding: 8px 12px; }
  </style>
</head>
<body>
  <div class="wrap">
    <h2>Generate Hall Ticket</h2>
    <p>Enter a register number (e.g., CSE001, ECE001, MECH001):</p>
    <form method="get" action="/qr">
      <input type="text" name="reg_no" placeholder="Register Number" required />
      <button type="submit">Show QR</button>
    </form>
  </div>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(INDEX_HTML)

def fetch_student_and_subjects(reg_no: str):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('SELECT reg_no, name, deg, branch, dob, sem, gender, semtime FROM students WHERE reg_no=?', (reg_no,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return None
    student = {
        'reg_no': row[0], 'name': row[1], 'deg': row[2], 'branch': row[3],
        'dob': row[4], 'sem': row[5], 'gender': row[6], 'semtime': row[7]
    }
    cur.execute('SELECT sem, date, session, code, name FROM subjects WHERE reg_no=?', (reg_no,))
    subjects = [
        {'sem': r[0], 'date': r[1], 'session': r[2], 'code': r[3], 'name': r[4]}
        for r in cur.fetchall()
    ]
    conn.close()
    return student, subjects

def find_docx_template() -> Path:
  # Pick first .docx in templates directory
  for name in os.listdir(TEMPLATES_DIR):
    if name.lower().endswith('.docx'):
      return Path(TEMPLATES_DIR) / name
  raise FileNotFoundError('No DOCX template found in templates/.')

def replace_in_doc(doc, repl: dict):
  # Replace in paragraphs
  for p in doc.paragraphs:
    for r in p.runs:
      for k, v in repl.items():
        if k in r.text:
          r.text = r.text.replace(k, str(v))
  # Replace in tables
  for t in doc.tables:
    for row in t.rows:
      for cell in row.cells:
        for p in cell.paragraphs:
          for r in p.runs:
            for k, v in repl.items():
              if k in r.text:
                r.text = r.text.replace(k, str(v))

def find_table_with(doc, placeholder: str):
  for t in doc.tables:
    for row in t.rows:
      for cell in row.cells:
        if placeholder in cell.text:
          return t
  return None

def remove_row(table, row):
  tbl = table._tbl
  tr = row._tr
  tbl.remove(tr)

def build_hall_ticket_pdf_bytes(reg_no: str, qr_url: str) -> tuple[bytes, str]:
  data = fetch_student_and_subjects(reg_no)
  if data is None:
    raise RuntimeError('Student not found')
  student, subjects = data
  template_path = find_docx_template()
  doc = Document(str(template_path))

  replacements = {
    '[name]': student['name'],
    '[regno]': student['reg_no'],
    '[Deg]': student['deg'],
    '[branch]': student['branch'],
    '[dob]': student['dob'],
    '[sem]': student['sem'],
    '[gen]': student['gender'],
    '[semtime]': student['semtime'],
  }
  replace_in_doc(doc, replacements)

  # Fill subjects table using placeholders [subcode] / [subname]
  subj_placeholder = '[subcode]'
  table = find_table_with(doc, subj_placeholder)
  if table:
    # find placeholder row
    placeholder_row = None
    for row in table.rows:
      row_text = ' | '.join([c.text for c in row.cells])
      if '[subcode]' in row_text or '[subname]' in row_text:
        placeholder_row = row
        break
    if placeholder_row:
      remove_row(table, placeholder_row)
    for s in subjects:
      cells = table.add_row().cells
      cells[0].text = s.get('sem', student['sem'])
      cells[1].text = s.get('date', '')
      cells[2].text = s.get('session', '')
      cells[3].text = s.get('code', '')
      cells[4].text = s.get('name', '')

  # Add QR if template has [QR_BOX]
  qr_table = find_table_with(doc, '[QR_BOX]')
  if qr_table:
    import qrcode
    qr_img = qrcode.make(qr_url)
    tmp_png = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    qr_img.save(tmp_png.name)
    for row in qr_table.rows:
      for cell in row.cells:
        if '[QR_BOX]' in cell.text:
          cell.text = ''
          run = cell.paragraphs[0].add_run()
          run.add_picture(tmp_png.name, width=Mm(30))
          break

  # Convert to PDF using temporary files
  tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
  tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
  doc.save(tmp_docx.name)
  # Try LibreOffice (soffice) headless conversion first (no Word required)
  soffice_paths = [
      r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
      r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
      r"/usr/bin/soffice",
      r"/usr/local/bin/soffice",
  ]
  soffice = next((p for p in soffice_paths if os.path.exists(p)), None)
  if soffice:
      out_dir = os.path.dirname(tmp_pdf.name)
      try:
          # Convert to PDF
          subprocess.run([
              soffice, '--headless', '--convert-to', 'pdf', '--outdir', out_dir, tmp_docx.name
          ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
          # LibreOffice outputs <filename>.pdf in out_dir; ensure we read the correct file
          generated_pdf = Path(tmp_docx.name).with_suffix('.pdf')
          if generated_pdf.exists():
              with open(generated_pdf, 'rb') as f:
                  pdf_bytes = f.read()
              result = (pdf_bytes, 'pdf')
              # Clean up generated PDF
              try:
                  os.unlink(generated_pdf)
              except Exception:
                  pass
          else:
              # Fallback to DOCX if conversion did not create the file
              with open(tmp_docx.name, 'rb') as fdoc:
                  result = (fdoc.read(), 'docx')
      except Exception:
          # Fallback to DOCX on conversion error
          with open(tmp_docx.name, 'rb') as fdoc:
              result = (fdoc.read(), 'docx')
  else:
      # No LibreOffice; return DOCX
      with open(tmp_docx.name, 'rb') as fdoc:
          result = (fdoc.read(), 'docx')
  # Cleanup temp files
  try:
    os.unlink(tmp_docx.name)
    os.unlink(tmp_pdf.name)
    if 'tmp_png' in locals():
      os.unlink(tmp_png.name)
  except Exception:
    pass
  return result

@app.route('/download/<reg_no>')
def download_on_demand(reg_no: str):
  try:
    ip = get_local_ip()
    qr_url = f"http://{ip}:5000/download/{reg_no}"
    data_bytes, kind = build_hall_ticket_pdf_bytes(reg_no, qr_url)
  except Exception as e:
    return f'Failed to generate file: {e}', 500
  resp = make_response(data_bytes)
  if kind == 'pdf':
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = f'attachment; filename={reg_no}_hallticket.pdf'
  else:
    resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    resp.headers['Content-Disposition'] = f'attachment; filename={reg_no}_hallticket.docx'
  return resp

@app.route('/qr')
def qr_form_redirect():
    reg_no = request.args.get('reg_no')
    if not reg_no:
        return redirect(url_for('index'))
    return redirect(url_for('qr_for', filename=reg_no))

@app.route('/qr/<path:filename>')
def qr_for(filename):
    import qrcode
    ip = get_local_ip()
    # Build absolute URL for download
    download_url = f"http://{ip}:5000/download/{filename}"
    # Create QR image
    qr = qrcode.QRCode(box_size=10, border=2)
    qr.add_data(download_url)
    qr.make(fit=True)
    img = qr.make_image(fill_color='black', back_color='white')
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    qr_b64 = base64.b64encode(buf.getvalue()).decode('ascii')
    return render_template_string(QR_TEMPLATE, qr_b64=qr_b64, download_url=download_url)

if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    app.run(host='0.0.0.0', port=5000)
